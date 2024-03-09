import streamlit as st
from docx import Document
import os

# Password protection
password = st.secrets["password"]
entered_password = st.text_input("Enter the password:", type="password")

if entered_password == "":
    entered_password = None

if entered_password != password:
    if entered_password is not None:
        st.error("Incorrect password. Please try again.")
    st.stop()

# Load the document
doc_path = 'Template_for_word_replace.docx'
doc = Document(doc_path)

# Function to replace text while keeping formatting
def replace_text_while_keeping_formatting(cell, search_text, replace_text):
    # First, extract all the runs in the cell's paragraphs
    runs = []
    for para in cell.paragraphs:
        for run in para.runs:
            runs.append(run)

    # Combine the text from all runs and replace the target text
    full_text = ''.join(run.text for run in runs)
    new_text = full_text.replace(search_text, replace_text)

    # Clear all the runs in the cell
    for para in cell.paragraphs:
        for run in para.runs:
            run.clear()

    # Repopulate the runs with the new text
    current_run = runs[0] if runs else cell.paragraphs[0].add_run()
    current_run.text = new_text

# Streamlit app
st.title("Document Modifier")

# Input for company name
company_name = st.text_input("Company Name:", "DOLLAR TREE IMPORT")

# Input for Address Line 1
address_line1 = st.text_input("Address Line 1:", "")

# Input for Address Line 2
address_line2 = st.text_input("Address Line 2:", "")

# Input for Address Line 3
address_line3 = st.text_input("Address Line 3:", "")

# Prepare address
address = f"{address_line1}\n{address_line2}\n{address_line3}"

# Input for factor
factor = st.number_input("Factor:", value=1.0, step=0.1)

# Button to generate modified document and final PDF
if st.button("Generate Documents"):
    # Load the document
    doc = Document(doc_path)

    # Define the replacement values
    unit_prices = {"U1": 0.44, "U2": 0.35, "U3": 0.35, "U4": 0.40, "U5": 0.54}
    quantity = {"Q1": 6588.0, "Q2": 6576.0, "Q3": 10020.0, "Q4": 9144.0, "Q5": 4248.0}

    # Calculate modified unit prices based on the factor
    unit_prices_modified = {key: value * factor for key, value in unit_prices.items()}

    # Replace placeholders in the document
    for paragraph in doc.paragraphs:
        if 'COMPANY_NAME' in paragraph.text:
            paragraph.text = paragraph.text.replace('COMPANY_NAME', company_name)
        if 'ADDRESS' in paragraph.text:
            paragraph.text = paragraph.text.replace('ADDRESS', address)
    # Initialize a flag to determine when to start summing the item totals
    start_summing = False
    # Initialize total_sum_rounded
    total_sum_rounded = 0.0
    # Iterate through the document's tables to replace placeholders
    for table in doc.tables:
        total_sum = 0.0
        for row in table.rows:
            for cell in row.cells:
                # Replace unit prices
                for key, value in unit_prices_modified.items():
                    value_rounded = round(value, 2)
                    replace_text_while_keeping_formatting(cell, key, f"\n\t{value_rounded}")
                # Replace quantities and calculate total item
                for key, value in quantity.items():
                    if key.startswith('Q'):
                        idx = int(key[1:])
                        total_item = value * unit_prices_modified[f'U{idx}']
                        if key=='Q1':
                            start_summing=True
                        if start_summing and key.startswith('Q'):
                            total_sum+=total_item
                        cell_idx = row.cells[-1]
                        replace_text_while_keeping_formatting(cell_idx, f'T{idx}', "{:.2f}".format(total_item))
    total=total_sum/5
    # Round off the total sum to two decimal places
    total_sum_rounded = round(total, 2)

    # Replace placeholder "T6" with the calculated total sum
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if 'T6' in cell.text:
                    replace_text_while_keeping_formatting(cell, 'T6', "{:.2f}".format(total_sum_rounded))

    # Save the modified document
    doc.save('modified_document.docx')

    # Convert the modified document to PDF
    pdf_path = os.path.abspath('final.pdf')
    doc_path_abs = os.path.abspath('modified_document.docx')
    os.system(f'libreoffice --headless --convert-to pdf {doc_path_abs} --outdir {os.path.dirname(pdf_path)}')

    st.success("Documents generated successfully!")
